"""
Resume builder views.

Handles CRUD operations for resumes and export functionality (PDF, DOCX).
"""

import io
import json
import logging

from django.http import HttpResponse, JsonResponse
from django.shortcuts import render, redirect, get_object_or_404

from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

from docx import Document

from .models import Resume

logger = logging.getLogger(__name__)


def home(request):
    """Display list of recent resumes."""
    resumes = Resume.objects.all()[:20]
    return render(request, "builder/home.html", {"resumes": resumes})


def delete_resume(request, resume_id: int):
    """Delete a resume. Returns JSON for AJAX or redirects to home."""
    resume = get_object_or_404(Resume, id=resume_id)
    try:
        resume.delete()
        if request.method == "POST":
            return JsonResponse({"status": "success"})
    except Exception as e:
        logger.error(f"Error deleting resume {resume_id}: {e}")
        if request.method == "POST":
            return JsonResponse({"status": "error"}, status=500)
    return redirect("home")


def resume_create(request):
    """Create a new resume with basic info."""
    if request.method == "POST":
        title = (request.POST.get("title") or "My Resume").strip()
        full_name = (request.POST.get("full_name") or "").strip()

        try:
            resume = Resume.objects.create(
                title=title,
                full_name=full_name,
            )
            return redirect("resume_edit", resume_id=resume.id)
        except Exception as e:
            logger.error(f"Error creating resume: {e}")
            return render(request, "builder/resume_basics.html", {"error": "Failed to create resume"})

    return render(request, "builder/resume_basics.html")


def _validate_resume_data(title, full_name, email, phone):
    """
    Validate resume basic information.
    
    Args:
        title: Resume title
        full_name: Full name
        email: Email address
        phone: Phone number
        
    Returns:
        Dictionary of errors (empty if valid)
    """
    errors = {}
    if not title:
        errors["title"] = "Title is required."
    if not full_name:
        errors["full_name"] = "Full name is required."
    if not email:
        errors["email"] = "Email is required."
    if not phone:
        errors["phone"] = "Phone number is required."
    return errors


def resume_edit(request, resume_id: int):
    """Edit resume details and structured sections."""
    resume = get_object_or_404(Resume, id=resume_id)
    errors = {}

    if request.method == "POST":
        # Gather basic info
        title = (request.POST.get("title") or resume.title).strip()
        full_name = (request.POST.get("full_name") or "").strip()
        email = (request.POST.get("email") or "").strip()
        phone = (request.POST.get("phone") or "").strip()

        # Validate
        errors = _validate_resume_data(title, full_name, email, phone)

        # If no errors, save the resume
        if not errors:
            try:
                resume.title = title
                resume.full_name = full_name
                resume.email = email
                resume.phone = phone
                resume.location = (request.POST.get("location") or "").strip()
                resume.linkedin = (request.POST.get("linkedin") or "").strip()
                resume.github = (request.POST.get("github") or "").strip()
                resume.summary = (request.POST.get("summary") or "").strip()

                # JSON sections
                resume.education = _safe_json_list(request.POST.get("education_json", "[]"))
                resume.experience = _safe_json_list(request.POST.get("experience_json", "[]"))
                resume.projects = _safe_json_list(request.POST.get("projects_json", "[]"))
                resume.skills = _safe_json_list(request.POST.get("skills_json", "[]"))

                resume.save()
                logger.info(f"Resume {resume_id} updated successfully")

                if request.POST.get("action") == "preview":
                    return redirect("resume_preview", resume_id=resume.id)

                return redirect("resume_edit", resume_id=resume.id)
            except Exception as e:
                logger.error(f"Error saving resume {resume_id}: {e}")
                errors["__all__"] = "Failed to save resume. Please try again."

    return render(
        request,
        "builder/resume_edit.html",
        {
            "resume": resume,
            "errors": errors,
            "education_json": json.dumps(resume.education or []),
            "experience_json": json.dumps(resume.experience or []),
            "projects_json": json.dumps(resume.projects or []),
            "skills_json": json.dumps(resume.skills or []),
        },
    )


def resume_preview(request, resume_id: int):
    """Preview resume in HTML."""
    resume = get_object_or_404(Resume, id=resume_id)
    return render(request, "builder/preview.html", {"resume": resume})


def _get_contact_line(resume):
    """Build contact line from email, phone, location."""
    parts = [resume.email, resume.phone, resume.location]
    return " | ".join([x for x in parts if x])


def _get_links_line(resume):
    """Build links line from LinkedIn and GitHub."""
    parts = [resume.linkedin, resume.github]
    return " | ".join([x for x in parts if x])


def _wrap_text(text, width=95):
    """
    Wrap text to specified width, splitting on word boundaries.
    
    Args:
        text: Text to wrap
        width: Maximum line width
        
    Returns:
        List of wrapped lines
    """
    words = (text or "").split()
    lines, current_line = [], []
    current_width = 0
    
    for word in words:
        word_len = len(word)
        space_len = 1 if current_line else 0
        
        if current_width + space_len + word_len <= width:
            current_line.append(word)
            current_width += space_len + word_len
        else:
            if current_line:
                lines.append(" ".join(current_line))
            current_line = [word]
            current_width = word_len
    
    if current_line:
        lines.append(" ".join(current_line))
    
    return lines


def export_pdf(request, resume_id: int):
    """Export resume as PDF."""
    resume = get_object_or_404(Resume, id=resume_id)

    try:
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=LETTER)
        width, height = LETTER
        y = height - 50

        def write_line(text, size=11, gap=16):
            """Write a line of text to the PDF."""
            nonlocal y
            c.setFont("Helvetica", size)
            # Truncate long text to avoid overflow
            display_text = (text or "")[:120]
            c.drawString(50, y, display_text)
            y -= gap
            if y < 60:
                c.showPage()
                y = height - 50

        # Header: Name
        write_line(resume.full_name or "Unnamed", size=16, gap=22)
        
        # Contact info
        contact = _get_contact_line(resume)
        if contact:
            write_line(contact, size=10, gap=14)

        # Links
        links = _get_links_line(resume)
        if links:
            write_line(links, size=10, gap=14)

        # Summary
        if resume.summary:
            write_line("")
            write_line("SUMMARY", size=12, gap=18)
            for chunk in _wrap_text(resume.summary, 95):
                write_line(chunk, size=10, gap=14)

        # Education
        if resume.education:
            write_line("")
            write_line("EDUCATION", size=12, gap=18)
            for edu in resume.education:
                degree = edu.get("degree", "")
                institution = edu.get("institution", "")
                dates = edu.get("dates", "")
                details = edu.get("details", "")
                
                title = " — ".join([x for x in [degree, institution] if x])
                write_line(f"{title}  {dates}".strip(), size=10, gap=14)
                
                if details:
                    for chunk in _wrap_text(details, 95):
                        write_line(f"  • {chunk}", size=10, gap=14)

        # Experience
        if resume.experience:
            write_line("")
            write_line("EXPERIENCE", size=12, gap=18)
            for exp in resume.experience:
                role = exp.get("role", "")
                company = exp.get("company", "")
                dates = exp.get("dates", "")
                
                title = " — ".join([x for x in [role, company] if x])
                write_line(f"{title}  {dates}".strip(), size=10, gap=14)
                
                for bullet in (exp.get("bullets") or []):
                    for chunk in _wrap_text(bullet, 95):
                        write_line(f"  • {chunk}", size=10, gap=14)

        # Projects
        if resume.projects:
            write_line("")
            write_line("PROJECTS", size=12, gap=18)
            for proj in resume.projects:
                name = proj.get("name", "Project")
                tech = proj.get("tech", "")
                
                write_line(f"{name}  {tech}".strip(), size=10, gap=14)
                
                for bullet in (proj.get("bullets") or []):
                    for chunk in _wrap_text(bullet, 95):
                        write_line(f"  • {chunk}", size=10, gap=14)

        # Skills
        if resume.skills:
            write_line("")
            write_line("SKILLS", size=12, gap=18)
            skills_text = ", ".join([s for s in resume.skills if isinstance(s, str)])
            write_line(skills_text[:110], size=10, gap=14)

        c.save()
        buffer.seek(0)
        filename = f"{(resume.full_name or 'resume').replace(' ', '_')}.pdf"

        logger.info(f"PDF exported for resume {resume_id}")

        return HttpResponse(
            buffer.getvalue(),
            content_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        logger.error(f"Error exporting PDF for resume {resume_id}: {e}")
        return HttpResponse("Error generating PDF", status=500)


def export_docx(request, resume_id: int):
    """Export resume as DOCX."""
    resume = get_object_or_404(Resume, id=resume_id)

    try:
        doc = Document()
        doc.add_heading(resume.full_name or "Unnamed", level=0)

        # Contact info
        contact = _get_contact_line(resume)
        if contact:
            doc.add_paragraph(contact)

        # Links
        links = _get_links_line(resume)
        if links:
            doc.add_paragraph(links)

        # Summary
        if resume.summary:
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(resume.summary)

        # Education
        if resume.education:
            doc.add_heading("Education", level=1)
            for edu in resume.education:
                degree = edu.get("degree", "")
                institution = edu.get("institution", "")
                dates = edu.get("dates", "")
                details = edu.get("details", "")
                
                title = " — ".join([x for x in [degree, institution, dates] if x])
                doc.add_paragraph(title, style="List Bullet")
                
                if details:
                    doc.add_paragraph(details)

        # Experience
        if resume.experience:
            doc.add_heading("Experience", level=1)
            for exp in resume.experience:
                role = exp.get("role", "")
                company = exp.get("company", "")
                dates = exp.get("dates", "")
                
                title = " — ".join([x for x in [role, company, dates] if x])
                doc.add_paragraph(title, style="List Bullet")
                
                for bullet in (exp.get("bullets") or []):
                    doc.add_paragraph(bullet, style="List Bullet 2")

        # Projects
        if resume.projects:
            doc.add_heading("Projects", level=1)
            for proj in resume.projects:
                name = proj.get("name", "")
                tech = proj.get("tech", "")
                
                title = " — ".join([x for x in [name, tech] if x])
                doc.add_paragraph(title, style="List Bullet")
                
                for bullet in (proj.get("bullets") or []):
                    doc.add_paragraph(bullet, style="List Bullet 2")

        # Skills
        if resume.skills:
            doc.add_heading("Skills", level=1)
            skills_text = ", ".join([s for s in resume.skills if isinstance(s, str)])
            doc.add_paragraph(skills_text)

        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        filename = f"{(resume.full_name or 'resume').replace(' ', '_')}.docx"

        logger.info(f"DOCX exported for resume {resume_id}")

        return HttpResponse(
            f.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        logger.error(f"Error exporting DOCX for resume {resume_id}: {e}")
        return HttpResponse("Error generating DOCX", status=500)


def _safe_json_list(raw: str):
    """
    Safely parse JSON list from string.
    
    Args:
        raw: Raw JSON string
        
    Returns:
        List if valid JSON, empty list otherwise
    """
    try:
        val = json.loads(raw or "[]")
        return val if isinstance(val, list) else []
    except (json.JSONDecodeError, TypeError):
        logger.warning(f"Failed to parse JSON: {raw}")
        return []

